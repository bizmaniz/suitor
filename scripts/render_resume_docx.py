#!/usr/bin/env python3
"""Render a markdown resume/letter draft to DOCX.

Supports: # / ## / ### headings, "- " bullets, **bold**, and [text](url)
markdown links rendered as real Word hyperlinks (blue + underlined).

Usage: render_resume_docx.py input.md output.docx
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt

LINK_SPLIT = re.compile(r"(\[[^\]]+\]\([^)]+\))")
LINK_MATCH = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")
BOLD_SPLIT = re.compile(r"(\*\*[^*]+\*\*)")
BULLET = re.compile(r"^[-*•]\s+")


def add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    props.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(underline)
    run.append(props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bold_runs(paragraph, text: str) -> None:
    for part in BOLD_SPLIT.split(text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part:
            paragraph.add_run(part)


def add_runs(paragraph, text: str) -> None:
    for part in LINK_SPLIT.split(text):
        match = LINK_MATCH.match(part)
        if match:
            add_hyperlink(paragraph, match.group(2), match.group(1))
        elif part:
            add_bold_runs(paragraph, part)


def main() -> int:
    if len(sys.argv) != 3:
        raise SystemExit("Usage: render_resume_docx.py input.md output.docx")
    source = Path(sys.argv[1]).read_text(encoding="utf-8")
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    for raw in source.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif BULLET.match(line):
            add_runs(doc.add_paragraph(style="List Bullet"), BULLET.sub("", line))
        else:
            add_runs(doc.add_paragraph(), line)

    doc.save(sys.argv[2])
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
